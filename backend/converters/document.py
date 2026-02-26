"""Conversor de documentos e dados tabulares."""
import json
from pathlib import Path


def convert(input_path: str, output_path: str, target_format: str) -> None:
    input_ext = Path(input_path).suffix.lstrip(".").lower()
    target_format = target_format.lower()

    # Dados tabulares
    if input_ext in ("csv", "json", "xlsx", "xls"):
        _convert_tabular(input_path, output_path, input_ext, target_format)
    elif input_ext == "pdf":
        _pdf_convert(input_path, output_path, target_format)
    elif input_ext == "docx":
        _docx_convert(input_path, output_path, target_format)
    elif input_ext in ("txt", "md", "html"):
        _text_convert(input_path, output_path, input_ext, target_format)
    else:
        raise ValueError(f"Conversor de documento não suporta: {input_ext}")


# --- Dados tabulares ---

def _convert_tabular(input_path, output_path, input_ext, target_format):
    import pandas as pd

    if input_ext == "csv":
        df = pd.read_csv(input_path)
    elif input_ext == "json":
        df = pd.read_json(input_path)
    elif input_ext in ("xlsx", "xls"):
        df = pd.read_excel(input_path)
    else:
        raise ValueError(f"Formato tabular desconhecido: {input_ext}")

    if target_format == "csv":
        df.to_csv(output_path, index=False)
    elif target_format == "json":
        df.to_json(output_path, orient="records", indent=2, force_ascii=False)
    elif target_format == "xlsx":
        df.to_excel(output_path, index=False)
    else:
        raise ValueError(f"Saída tabular não suportada: {target_format}")


# --- PDF ---

def _pdf_convert(input_path, output_path, target_format):
    import fitz  # PyMuPDF

    doc = fitz.open(input_path)

    if target_format == "txt":
        text = "\n\n".join(page.get_text() for page in doc)
        Path(output_path).write_text(text, encoding="utf-8")

    elif target_format == "html":
        html_parts = [
            "<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>"
        ]
        for i, page in enumerate(doc):
            html_parts.append(f"<div class='page' id='page-{i+1}'>")
            html_parts.append(page.get_text("html"))
            html_parts.append("</div>")
        html_parts.append("</body></html>")
        Path(output_path).write_text("\n".join(html_parts), encoding="utf-8")

    elif target_format == "png":
        # Renderiza todas as páginas em um único PNG vertical
        from PIL import Image
        import io

        images = []
        for page in doc:
            mat = fitz.Matrix(2, 2)  # 2x zoom = ~144dpi
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            images.append(img)

        if not images:
            raise RuntimeError("PDF sem páginas")

        total_height = sum(img.height for img in images)
        max_width = max(img.width for img in images)
        combined = Image.new("RGB", (max_width, total_height), "white")
        y_offset = 0
        for img in images:
            combined.paste(img, (0, y_offset))
            y_offset += img.height
        combined.save(output_path, format="PNG")

    else:
        raise ValueError(f"PDF não suporta saída: {target_format}")

    doc.close()


# --- DOCX ---

def _docx_convert(input_path, output_path, target_format):
    from docx import Document

    doc = Document(input_path)

    if target_format == "txt":
        text = "\n".join(p.text for p in doc.paragraphs)
        Path(output_path).write_text(text, encoding="utf-8")

    elif target_format == "html":
        lines = ["<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>"]
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading"):
                level = p.style.name.split()[-1]
                lines.append(f"<h{level}>{p.text}</h{level}>")
            else:
                lines.append(f"<p>{p.text}</p>")
        lines.append("</body></html>")
        Path(output_path).write_text("\n".join(lines), encoding="utf-8")

    elif target_format == "pdf":
        import weasyprint
        # Converte primeiro para HTML, depois para PDF
        html_tmp = output_path + ".tmp.html"
        _docx_convert(input_path, html_tmp, "html")
        weasyprint.HTML(filename=html_tmp).write_pdf(output_path)
        Path(html_tmp).unlink(missing_ok=True)

    else:
        raise ValueError(f"DOCX não suporta saída: {target_format}")


# --- Texto simples / Markdown / HTML ---

def _text_convert(input_path, output_path, input_ext, target_format):
    text = Path(input_path).read_text(encoding="utf-8")

    if input_ext == "md":
        if target_format == "html":
            import markdown as md_lib
            html_body = md_lib.markdown(text)
            html = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{html_body}</body></html>"
            Path(output_path).write_text(html, encoding="utf-8")
        elif target_format == "txt":
            # Remove marcações Markdown simples
            import re
            plain = re.sub(r"[#*`_~>\[\]()]+", "", text)
            Path(output_path).write_text(plain, encoding="utf-8")
        elif target_format == "pdf":
            import markdown as md_lib
            import weasyprint
            html_body = md_lib.markdown(text)
            html = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{html_body}</body></html>"
            weasyprint.HTML(string=html).write_pdf(output_path)
        else:
            raise ValueError(f"MD não suporta saída: {target_format}")

    elif input_ext == "txt":
        if target_format == "html":
            escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body><pre>{escaped}</pre></body></html>"
            Path(output_path).write_text(html, encoding="utf-8")
        elif target_format == "md":
            Path(output_path).write_text(text, encoding="utf-8")
        elif target_format == "pdf":
            import weasyprint
            escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html = f"<html><body><pre style='font-family:monospace'>{escaped}</pre></body></html>"
            weasyprint.HTML(string=html).write_pdf(output_path)
        elif target_format == "docx":
            from docx import Document
            doc = Document()
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.save(output_path)
        else:
            raise ValueError(f"TXT não suporta saída: {target_format}")

    elif input_ext == "html":
        if target_format == "txt":
            import re
            plain = re.sub(r"<[^>]+>", "", text)
            Path(output_path).write_text(plain, encoding="utf-8")
        elif target_format == "md":
            # Conversão básica HTML → Markdown
            import re
            result = text
            result = re.sub(r"<h([1-6])>(.*?)</h\1>", lambda m: "#" * int(m.group(1)) + " " + m.group(2), result, flags=re.IGNORECASE)
            result = re.sub(r"<b>(.*?)</b>", r"**\1**", result, flags=re.IGNORECASE)
            result = re.sub(r"<i>(.*?)</i>", r"*\1*", result, flags=re.IGNORECASE)
            result = re.sub(r"<p>(.*?)</p>", r"\1\n", result, flags=re.IGNORECASE | re.DOTALL)
            result = re.sub(r"<[^>]+>", "", result)
            Path(output_path).write_text(result, encoding="utf-8")
        elif target_format == "pdf":
            import weasyprint
            weasyprint.HTML(filename=input_path).write_pdf(output_path)
        else:
            raise ValueError(f"HTML não suporta saída: {target_format}")
